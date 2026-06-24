from __future__ import annotations

"""Lightweight knowledge-base ingestion and hybrid retrieval.

This module intentionally starts simple: documents are saved locally, text is
chunked into SQLite, and retrieval combines keyword matching with character
n-gram similarity. It is enough for a deployable demo and leaves a clean
upgrade path to vector search later.
"""

import re
import uuid
from pathlib import Path

from fastapi import UploadFile

from app.database import (
    add_knowledge_chunks,
    create_knowledge_document,
    list_knowledge_chunks,
)
from app.schemas import ChatMessage, KnowledgeDocument, ToolResult

BASE_DIR = Path(__file__).resolve().parent.parent
KNOWLEDGE_DIR = BASE_DIR / "data" / "knowledge"
MAX_FILE_SIZE = 12 * 1024 * 1024
MAX_CONTEXT_CHARS = 2600
SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}
DOMAIN_SYNONYMS = {
    "学生证": ("校园卡", "证件", "学生卡"),
    "补办": ("重新办理", "重新办", "补领", "办理"),
    "遗失": ("丢了", "丢失", "遗失", "找不到"),
    "报修": ("维修", "故障", "坏了", "修理"),
    "空调": ("制冷", "不冷", "不制冷"),
    "宿舍": ("寝室", "公寓"),
}


async def ingest_knowledge_file(file: UploadFile) -> KnowledgeDocument:
    """Save one uploaded document, extract text, chunk it, and persist metadata."""
    raw = await file.read()
    if not raw:
        raise ValueError("知识库文件为空")
    if len(raw) > MAX_FILE_SIZE:
        raise ValueError("知识库文件不能超过 12MB")

    filename = Path(file.filename or "knowledge.txt").name
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("仅支持 txt、md、pdf、docx 知识库文件")

    KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid.uuid4().hex}{extension}"
    stored_path = KNOWLEDGE_DIR / stored_name
    stored_path.write_bytes(raw)

    text = extract_text(raw, extension)
    chunks = chunk_text(text)
    if not chunks:
        raise ValueError("没有从文件中提取到可检索文本")

    row = create_knowledge_document(
        filename=filename,
        content_type=file.content_type or "application/octet-stream",
        path=str(stored_path),
        chunk_count=len(chunks),
    )
    add_knowledge_chunks(row["id"], chunks)
    return KnowledgeDocument(
        id=row["id"],
        filename=row["filename"],
        content_type=row["content_type"],
        chunk_count=row["chunk_count"],
        created_at=row["created_at"],
    )


def extract_text(raw: bytes, extension: str) -> str:
    """Extract text from supported file types."""
    if extension in {".txt", ".md", ".markdown"}:
        return decode_text(raw)
    if extension == ".pdf":
        return extract_pdf_text(raw)
    if extension == ".docx":
        return extract_docx_text(raw)
    return ""


def decode_text(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def extract_pdf_text(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("解析 PDF 需要安装 pypdf，请执行 pip install pypdf") from exc

    from io import BytesIO

    reader = PdfReader(BytesIO(raw))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("解析 Word 需要安装 python-docx，请执行 pip install python-docx") from exc

    from io import BytesIO

    document = Document(BytesIO(raw))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def chunk_text(text: str, chunk_size: int = 700, overlap: int = 120) -> list[str]:
    """Split text into overlapping chunks that fit comfortably in prompts."""
    cleaned = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not cleaned:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(cleaned):
        end = min(len(cleaned), start + chunk_size)
        chunk = cleaned[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(cleaned):
            break
        start = max(0, end - overlap)
    return chunks


def retrieve_knowledge_result(messages: list[ChatMessage]) -> ToolResult | None:
    """Search document chunks for the latest user query and return prompt context."""
    query = next((message.content for message in reversed(messages) if message.role == "user"), "")
    hits = search_knowledge(query)
    if not hits:
        return None

    parts: list[str] = []
    total_chars = 0
    for hit in hits:
        snippet = f"来源：{hit['filename']}，片段 {hit['chunk_index'] + 1}\n{hit['content']}"
        if total_chars + len(snippet) > MAX_CONTEXT_CHARS:
            break
        parts.append(snippet)
        total_chars += len(snippet)

    if not parts:
        return None

    return ToolResult(
        name="knowledge_base",
        title="知识库",
        content=(
            "以下是从本地知识库检索到的资料片段。回答校园办事、报修、制度类问题时，"
            "优先依据这些资料；资料不足时要说明缺少依据。\n\n"
            + "\n\n---\n\n".join(parts)
        ),
    )


def search_knowledge(query: str, limit: int = 4) -> list[dict[str, object]]:
    """Rank chunks with a lightweight hybrid retrieval score.

    The score combines:
    - keyword overlap for precise terms such as names, phone numbers, places
    - filename/title boost for document-level intent
    - character n-gram similarity for Chinese paraphrases
    """
    terms = tokenize(query)
    query_grams = char_ngrams(query)
    if not terms and not query_grams:
        return []

    scored: list[dict[str, object]] = []
    for row in list_knowledge_chunks():
        content = row["content"]
        content_lower = content.lower()
        filename_lower = row["filename"].lower()
        keyword_score = 0.0
        title_score = 0.0
        for term in terms:
            keyword_score += content_lower.count(term)
            if term in filename_lower:
                title_score += 2.0

        semantic_score = jaccard_similarity(query_grams, char_ngrams(content))
        score = keyword_score * 1.0 + title_score * 1.2 + semantic_score * 8.0
        if score > 0:
            scored.append(
                {
                    "score": score,
                    "keyword_score": round(keyword_score, 4),
                    "semantic_score": round(semantic_score, 4),
                    "title_score": round(title_score, 4),
                    "filename": row["filename"],
                    "chunk_index": row["chunk_index"],
                    "content": content,
                }
            )

    scored.sort(key=lambda item: item["score"], reverse=True)
    return scored[:limit]


def tokenize(text: str) -> list[str]:
    """Build mixed Chinese/English keywords for lightweight retrieval."""
    lower = text.lower()
    words = re.findall(r"[a-z0-9_]+", lower)
    chinese_terms = re.findall(r"[\u4e00-\u9fff]{2,}", lower)
    short_phrases: list[str] = []
    for term in chinese_terms:
        if len(term) <= 6:
            short_phrases.append(term)
        else:
            short_phrases.extend(term[index : index + 2] for index in range(len(term) - 1))
    expanded_terms = expand_domain_terms(lower)
    unique_terms = dict.fromkeys([*words, *short_phrases, *expanded_terms])
    return [term for term in unique_terms if len(term) >= 2]


def expand_domain_terms(text: str) -> list[str]:
    """Add small domain synonyms before a full embedding service is introduced."""
    expanded: list[str] = []
    for canonical, aliases in DOMAIN_SYNONYMS.items():
        if canonical in text or any(alias in text for alias in aliases):
            expanded.append(canonical)
            expanded.extend(aliases)
    return expanded


def char_ngrams(text: str, n: int = 2) -> set[str]:
    """Create character n-grams for small-scale semantic-ish matching."""
    normalized = re.sub(r"\s+", "", text.lower())
    if len(normalized) < n:
        return {normalized} if normalized else set()
    return {normalized[index : index + n] for index in range(len(normalized) - n + 1)}


def jaccard_similarity(left: set[str], right: set[str]) -> float:
    if not left or not right:
        return 0.0
    return len(left & right) / len(left | right)
